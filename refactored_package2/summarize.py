import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from gensim.summarization.summarizer import summarize

def generateBriefing(sectionsDictionary): 
    document = Document()

    document.add_heading('Earthquake Report', 0)
    
    document.add_heading("Buildings", level = 1) 
    document.add_paragraph(sectionsDictionary["Buildings"]) 
    
    document.add_heading("Resilience", level = 1) 
    document.add_paragraph(sectionsDictionary["Resilience"]) 
    
    document.add_heading("Infrastructure", level = 1) 
    document.add_paragraph(sectionsDictionary["Infrastructure"]) 
    
    document.add_page_break()
    
    document.save("briefing.docx") 

    print(sectionsDictionary) 
    
def get_summary(df, section, ratio=0.1):
    df = df[df["majority"] == section]
    text = " ".join(df["sentence"].tolist())
    text_clean = set(summarize(text, ratio=ratio).split("\n")) # remove duplicated sentences
    return " ".join(text_clean)